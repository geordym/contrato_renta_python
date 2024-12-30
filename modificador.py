import json
from docx import Document
import boto3

s3 = boto3.client('s3')

def lambda_handler(event, context):
    # Parámetros de entrada recibidos por el evento API Gateway
    params = event.get("body", {})
    
    if isinstance(params, str):
        params = json.loads(params)
    
     input_bucket = os.environ['INPUT_BUCKET']
    output_bucket = os.environ['OUTPUT_BUCKET']
    input_key = os.environ['INPUT_KEY']
    output_key = os.environ['OUTPUT_KEY']

    
    # Descargar el archivo de S3
    s3.download_file(input_bucket, input_key, input_path)
    
    # Cargar el documento de Word
    doc = Document(input_path)
    
    # Diccionario de variables para reemplazar
    placeholders = {
        "{{PRIMER_PARRAFO_ARRENDADOR_NOMBRE}}": params.get("PRIMER_PARRAFO_ARRENDADOR_NOMBRE", ""),
        "{{PRIMER_PARRAFO_ARRENDATARIO_NOMBRE}}": params.get("PRIMER_PARRAFO_ARRENDATARIO_NOMBRE", ""),
        "{{DECLARACIONES_ARRENDADOR_NACIONALIDAD}}": params.get("DECLARACIONES_ARRENDADOR_NACIONALIDAD", ""),
        "{{DECLARACIONES_ARRENDADOR_TIPO_DOCUMENTO}}": params.get("DECLARACIONES_ARRENDADOR_TIPO_DOCUMENTO", ""),
        "{{DECLARACIONES_ARRENDADOR_NUMERO_DOCUMENTO}}": params.get("DECLARACIONES_ARRENDADOR_NUMERO_DOCUMENTO", ""),
        "{{DECLARACIONES_ARRENDADOR_DOMICILIO_FISCAL_DIRECCION}}": params.get("DECLARACIONES_ARRENDADOR_DOMICILIO_FISCAL_DIRECCION", ""),
        "{{DECLARACIONES_ARRENDADOR_DOMICILIO_PARTICULAR_DIRECCION}}": params.get("DECLARACIONES_ARRENDADOR_DOMICILIO_PARTICULAR_DIRECCION", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_MARCA}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_MARCA", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_MODELO}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_MODELO", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_COLOR}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_COLOR", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_TRANSMISION}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_TRANSMISION", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_PUERTAS}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_PUERTAS", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_PASAJEROS}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_PASAJEROS", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_PLACAS_CIRCULACION}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_PLACAS_CIRCULACION", ""),
        "{{DECLARACIONES_ARRENDADOR_VEHICULO_ESTADO_GEOGRAFICO}}": params.get("DECLARACIONES_ARRENDADOR_VEHICULO_ESTADO_GEOGRAFICO", ""),
        
        "{{DECLARACIONES_ARRENDATARIO_NACIONALIDAD}}": params.get("DECLARACIONES_ARRENDATARIO_NACIONALIDAD", ""),
        "{{DECLARACIONES_ARRENDATARIO_TIPO_DOCUMENTO}}": params.get("DECLARACIONES_ARRENDATARIO_TIPO_DOCUMENTO", ""),
        "{{DECLARACIONES_ARRENDATARIO_NUMERO_DOCUMENTO}}": params.get("DECLARACIONES_ARRENDATARIO_NUMERO_DOCUMENTO", ""),
        "{{DECLARACIONES_ARRENDATARIO_DOMICILIO}}": params.get("DECLARACIONES_ARRENDATARIO_DOMICILIO", ""),
        "{{DECLARACIONES_ARRENDATARIO_CIUDAD}}": params.get("DECLARACIONES_ARRENDATARIO_CIUDAD", ""),
        "{{DECLARACIONES_ARRENDATARIO_MUNICIPIO}}": params.get("DECLARACIONES_ARRENDATARIO_MUNICIPIO", ""),
        "{{DECLARACIONES_ARRENDATARIO_LICENCIA_EXPEDIDA_ESTADO}}": params.get("DECLARACIONES_ARRENDATARIO_LICENCIA_EXPEDIDA_ESTADO", ""),
        "{{DECLARACIONES_ARRENDATARIO_LICENCIA_NUMERO_IDENTIFICACION}}": params.get("DECLARACIONES_ARRENDATARIO_LICENCIA_NUMERO_IDENTIFICACION", ""),
        
        "{{CLAUSULAS_SEGUNDA_ARRENDAMIENTO_DURACION_DIAS}}": params.get("CLAUSULAS_SEGUNDA_ARRENDAMIENTO_DURACION_DIAS", ""),
        "{{CLAUSULAS_SEGUNDA_ARRENDAMIENTO_FECHA_INICIO}}": params.get("CLAUSULAS_SEGUNDA_ARRENDAMIENTO_FECHA_INICIO", ""),
        "{{CLAUSULAS_SEGUNDA_ARRENDAMIENTO_FECHA_FIN}}": params.get("CLAUSULAS_SEGUNDA_ARRENDAMIENTO_FECHA_FIN", ""),
        "{{CLÁUSULAS_SEGUNDA_ARRENDAMIENTO_VIGENCIA_DIAS}}": params.get("CLÁUSULAS_SEGUNDA_ARRENDAMIENTO_VIGENCIA_DIAS", ""),
        
        "{{CLAUSULAS_TERCERA_ARRENDAMIENTO_PAGO_CONCEPTO}}": params.get("CLAUSULAS_TERCERA_ARRENDAMIENTO_PAGO_CONCEPTO", ""),
        "{{CLAUSULAS_TERCERA_ARRENDAMIENTO_PAGO_MONTO}}": params.get("CLAUSULAS_TERCERA_ARRENDAMIENTO_PAGO_MONTO", ""),
        
        "{{CLAUSULAS_QUINTA_ARRENDAMIENTO_VEHICULO_USO}}": params.get("CLAUSULAS_QUINTA_ARRENDAMIENTO_VEHICULO_USO", ""),
        "{{CLAUSULAS_QUINTA_ARRENDAMIENTO_DESTINO}}": params.get("CLAUSULAS_QUINTA_ARRENDAMIENTO_DESTINO", ""),
        "{{CLAUSULAS_QUINTA_ARRENDAMIENTO_VEHICULO_CONDUCTOR}}": params.get("CLAUSULAS_QUINTA_ARRENDAMIENTO_VEHICULO_CONDUCTOR", ""),
        "{{CLAUSULAS_QUINTA_ARRENDAMIENTO_ARRENDATARIO_LICENCIA_NUMERO_IDENTIFICACION}}": params.get("CLAUSULAS_QUINTA_ARRENDAMIENTO_ARRENDATARIO_LICENCIA_NUMERO_IDENTIFICACION", ""),
        
        "{{CLAUSULAS_DECIMA_TERCERA_UBICACION_FIRMA}}": params.get("CLAUSULAS_DECIMA_TERCERA_UBICACION_FIRMA", ""),
        "{{CLAUSULAS_DECIMA_TERCERA_FECHA_FIRMA}}": params.get("CLAUSULAS_DECIMA_TERCERA_FECHA_FIRMA", ""),
        
        "{{FIRMAS_ZONA_ARRENDADOR_NOMBRE}}": params.get("FIRMAS_ZONA_ARRENDADOR_NOMBRE", ""),
        "{{FIRMAS_ZONA_ARRENDATARIO_NOMBRE}}": params.get("FIRMAS_ZONA_ARRENDATARIO_NOMBRE", ""),
        
        "{{PARRAFO_FINAL_ARRENDATARIO_DOCUMENTOS_TEXTO}}": params.get("PARRAFO_FINAL_ARRENDATARIO_DOCUMENTOS_TEXTO", "")
    }

    # Reemplazar las variables en los párrafos
    for para in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, value)
    
    # Reemplazar también en las celdas de las tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    
    # Guardar el documento modificado en un archivo temporal
    doc.save(output_path)
    
    # Subir el archivo modificado a S3 (opcional)
    output_bucket = 'your-bucket-name'  # Cambia esto con tu bucket de salida
    output_key = 'contrato_arrendamiento_modificado.docx'  # Ruta de archivo en S3
    s3.upload_file(output_path, output_bucket, output_key)

    # Retornar una respuesta JSON
    return {
        'statusCode': 200,
        'body': json.dumps({
            'message': 'Documento generado exitosamente',
            'output_file': output_key  # Devuelve el nombre del archivo subido a S3
        })
    }
